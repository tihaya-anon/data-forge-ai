"""
Document Processing Module for DataForge AI

This module provides functionality for:
- Document parsing (PDF, DOCX, TXT)
- Text chunking (fixed length/semantic chunking)
- Metadata extraction
- Document deduplication
"""

from typing import List, Dict, Any, Optional, Iterator
import hashlib
import os
from pathlib import Path


class Document:
    """Represents a processed document with content and metadata."""
    
    def __init__(self, content: str, metadata: Dict[str, Any] = None, doc_id: str = None):
        self.content = content
        self.metadata = metadata or {}
        self.id = doc_id or self._generate_doc_id(content)
    
    def _generate_doc_id(self, content: str) -> str:
        """Generate a unique document ID based on content hash."""
        return hashlib.sha256(content.encode()).hexdigest()


class DocumentProcessor:
    """Main class for processing documents."""
    
    def __init__(self):
        self.parsers = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.txt': self._parse_txt,
        }
    
    def parse_document(self, file_path: str) -> Document:
        """
        Parse a document based on its file extension.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Parsed Document object
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext not in self.parsers:
            raise ValueError(f"Unsupported file format: {ext}")
        
        # Extract content and basic metadata
        content = self.parsers[ext](file_path)
        metadata = {
            'source_file': str(path.name),
            'file_path': str(path.absolute()),
            'file_size': path.stat().st_size,
            'file_type': ext,
            'created_at': path.stat().st_ctime,
            'modified_at': path.stat().st_mtime,
        }
        
        return Document(content=content, metadata=metadata)
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file using PyPDF2 or similar library."""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 is required to parse PDF files. Install it with: pip install PyPDF2")
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            content = []
            for page in reader.pages:
                content.append(page.extract_text())
            return "\n".join(content)
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file using python-docx library."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required to parse DOCX files. Install it with: pip install python-docx")
        
        doc = DocxDocument(file_path)
        content = []
        for para in doc.paragraphs:
            content.append(para.text)
        return "\n".join(content)
    
    def _parse_txt(self, file_path: str) -> str:
        """Parse plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def chunk_document(
        self, 
        document: Document, 
        chunk_size: int = 512, 
        chunk_overlap: int = 50,
        method: str = 'fixed'
    ) -> List[Document]:
        """
        Split a document into chunks using either fixed-length or semantic chunking.
        
        Args:
            document: Document to chunk
            chunk_size: Size of each chunk
            chunk_overlap: Overlap between chunks
            method: 'fixed' or 'semantic'
            
        Returns:
            List of chunked documents
        """
        if method == 'fixed':
            return self._fixed_chunking(document, chunk_size, chunk_overlap)
        elif method == 'semantic':
            return self._semantic_chunking(document, chunk_size, chunk_overlap)
        else:
            raise ValueError(f"Unknown chunking method: {method}")
    
    def _fixed_chunking(
        self, 
        document: Document, 
        chunk_size: int, 
        chunk_overlap: int
    ) -> List[Document]:
        """Perform fixed-length chunking."""
        text = document.content
        chunks = []
        
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk_text = text[start:end]
            
            # Create chunk metadata
            chunk_metadata = document.metadata.copy()
            chunk_metadata['chunk_index'] = len(chunks)
            chunk_metadata['chunk_start'] = start
            chunk_metadata['chunk_end'] = end
            
            chunk_doc = Document(
                content=chunk_text,
                metadata=chunk_metadata
            )
            chunks.append(chunk_doc)
            
            # Move start position by chunk_size minus overlap
            start += chunk_size - chunk_overlap
        
        return chunks
    
    def _semantic_chunking(
        self, 
        document: Document, 
        chunk_size: int, 
        chunk_overlap: int
    ) -> List[Document]:
        """
        Perform semantic chunking by splitting on sentence boundaries.
        This is a simplified version; a full implementation would use NLP techniques.
        """
        import re
        
        # Split text by sentence boundaries
        sentences = re.split(r'[.!?]+\s+', document.content)
        
        chunks = []
        current_chunk = ""
        current_length = 0
        chunk_index = 0
        
        for sentence in sentences:
            # Check if adding this sentence would exceed chunk size
            if current_length + len(sentence) > chunk_size and current_chunk:
                # Finalize current chunk
                chunk_metadata = document.metadata.copy()
                chunk_metadata['chunk_index'] = chunk_index
                chunk_metadata['chunk_method'] = 'semantic'
                
                chunk_doc = Document(
                    content=current_chunk.strip(),
                    metadata=chunk_metadata
                )
                chunks.append(chunk_doc)
                
                # Start a new chunk with overlap
                if chunk_overlap > 0:
                    # Find the last few words in the current chunk to use as overlap
                    words = current_chunk.split()
                    overlap_words = words[-(chunk_overlap // 5):]  # Approximate overlap
                    current_chunk = " ".join(overlap_words) + " " + sentence
                else:
                    current_chunk = sentence
                
                current_length = len(current_chunk)
                chunk_index += 1
            else:
                current_chunk += " " + sentence
                current_length += len(sentence)
        
        # Add the final chunk if it has content
        if current_chunk.strip():
            chunk_metadata = document.metadata.copy()
            chunk_metadata['chunk_index'] = chunk_index
            chunk_metadata['chunk_method'] = 'semantic'
            
            chunk_doc = Document(
                content=current_chunk.strip(),
                metadata=chunk_metadata
            )
            chunks.append(chunk_doc)
        
        return chunks
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing extracted metadata
        """
        path = Path(file_path)
        stat = path.stat()
        
        metadata = {
            'source_file': str(path.name),
            'file_path': str(path.absolute()),
            'file_size': stat.st_size,
            'file_type': path.suffix.lower(),
            'created_at': stat.st_ctime,
            'modified_at': stat.st_mtime,
        }
        
        # Add format-specific metadata if available
        try:
            if path.suffix.lower() == '.pdf':
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    if reader.metadata:
                        for key, value in reader.metadata.items():
                            metadata[key.replace('/', '')] = value  # Remove '/' prefix from keys
                    metadata['page_count'] = len(reader.pages)
        except Exception:
            # If PDF metadata extraction fails, continue without it
            pass
        
        return metadata
    
    def deduplicate_documents(self, documents: List[Document], threshold: float = 0.95) -> List[Document]:
        """
        Remove duplicate documents based on content similarity.
        
        Args:
            documents: List of documents to deduplicate
            threshold: Similarity threshold for considering documents as duplicates (0-1)
            
        Returns:
            List of deduplicated documents
        """
        if not documents:
            return []
        
        # Simple deduplication based on content hash
        unique_docs = []
        seen_hashes = set()
        
        for doc in documents:
            content_hash = hashlib.sha256(doc.content.encode()).hexdigest()
            
            if content_hash not in seen_hashes:
                seen_hashes.add(content_hash)
                unique_docs.append(doc)
        
        # If we need more sophisticated deduplication based on similarity,
        # we could implement fuzzy matching here
        return unique_docs


# Convenience functions for common operations

def load_and_process_document(
    file_path: str, 
    chunk_size: int = 512, 
    chunk_overlap: int = 50,
    chunk_method: str = 'fixed'
) -> List[Document]:
    """
    Load, parse, and chunk a document in one step.
    
    Args:
        file_path: Path to the document
        chunk_size: Size of each chunk
        chunk_overlap: Overlap between chunks
        chunk_method: Method for chunking ('fixed' or 'semantic')
        
    Returns:
        List of chunked documents
    """
    processor = DocumentProcessor()
    document = processor.parse_document(file_path)
    chunks = processor.chunk_document(
        document, 
        chunk_size=chunk_size, 
        chunk_overlap=chunk_overlap,
        method=chunk_method
    )
    return chunks